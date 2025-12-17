import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import os
from datetime import datetime
import io
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import plotly.io as pio  # 保留但不用于图片导出

# -------------------------- 全局配置 --------------------------
st.set_page_config(
    page_title="企业数字化转型指数分析平台",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 初始化会话状态
if 'data_loaded' not in st.session_state:
    st.session_state.data_loaded = False
if 'merged_data' not in st.session_state:
    st.session_state.merged_data = None
if 'current_report_data' not in st.session_state:
    st.session_state.current_report_data = None
if 'chart_figs' not in st.session_state:
    st.session_state.chart_figs = {}

# -------------------------- 精准行业映射表 --------------------------
CODE_INDUSTRY_MAP = {
    "000001": "货币金融服务", "601318": "保险", "600036": "货币金融服务",
    "000002": "房地产业", "002594": "计算机应用", "600519": "酒类",
    "000858": "酒类", "002594": "汽车制造业", "601633": "汽车制造业",
    "600011": "电力", "600027": "电力"
}

NAME_INDUSTRY_MAP = {
    "深发展A": "货币金融服务", "平安银行": "货币金融服务", "中国平安": "保险",
    "贵州茅台": "酒类", "五粮液": "酒类", "伊利股份": "乳制品",
    "比亚迪": "汽车制造业", "长城汽车": "汽车制造业",
    "零七股份": "住宿业", "全新好": "住宿业", "*ST全新": "住宿业"
}

# -------------------------- 路径配置 --------------------------
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, "年报下载")
WORDFREQ_FILE = os.path.join(DATA_DIR, "词频数据.xlsx")
INDUSTRY_FILE = os.path.join(DATA_DIR, "最终数据dta格式-上市公司年度行业代码至2021.xlsx")

# 技术维度列
TECH_DIM_COLS = [
    '人工智能', '区块链', '大数据', '云计算', '物联网',
    '数字技术应用', '企业数字化', '数字运营', '数字安全',
    '5G通信', '数字平台', '数字人才'
]

# 自定义配色
COLOR_PALETTE = {
    'primary': '#2E86AB',
    'secondary': '#E63946',
    'accent': '#F1C40F',
    'neutral': '#A8DADC',
    'dark': '#1D3557'
}

# -------------------------- 核心工具函数 --------------------------
@st.cache_data(ttl=3600)
def load_data():
    """优化数据加载：仅加载必要列，减少内存占用"""
    try:
        if not os.path.exists(DATA_DIR):
            os.makedirs(DATA_DIR, exist_ok=True)
            return None, f"❌ 数据目录不存在，已自动创建：{DATA_DIR}\n请将词频数据和行业数据放入该目录后重试"

        if not os.path.exists(WORDFREQ_FILE):
            return None, f"❌ 词频文件不存在：{WORDFREQ_FILE}\n请确认文件路径是否正确"

        # 仅加载必要列
        wordfreq_df = pd.read_excel(
            WORDFREQ_FILE,
            engine='openpyxl',
            usecols=['股票代码', '年份', '企业名称', '总词频'] + TECH_DIM_COLS,
            dtype={'股票代码': str, '年份': int, '企业名称': str, '总词频': int}
        )
        wordfreq_df = wordfreq_df[wordfreq_df['年份'] >= 2010]  # 过滤旧数据

        industry_df = None
        if os.path.exists(INDUSTRY_FILE):
            industry_df = pd.read_excel(
                INDUSTRY_FILE,
                engine='openpyxl',
                usecols=['股票代码全称', '年度', '行业名称'],
                dtype={'股票代码全称': str, '年度': int, '行业名称': str}
            )
            industry_df.rename(columns={'股票代码全称': '股票代码', '年度': '年份', '行业名称': '申万行业名称'}, inplace=True)
            industry_df = industry_df[industry_df['年份'] >= 2010]
            merged_df = pd.merge(wordfreq_df, industry_df, on=['股票代码', '年份'], how='left')
        else:
            merged_df = wordfreq_df.copy()
            merged_df['申万行业名称'] = '未匹配行业'
            st.warning(f"⚠️ 行业数据文件不存在：{INDUSTRY_FILE}\n将使用精准映射表补全行业信息")

        # 精准行业匹配
        def get_industry(row):
            if row['股票代码'] in CODE_INDUSTRY_MAP:
                return CODE_INDUSTRY_MAP[row['股票代码']]
            elif row['企业名称'] in NAME_INDUSTRY_MAP:
                return NAME_INDUSTRY_MAP[row['企业名称']]
            else:
                return row['申万行业名称'] if pd.notna(row['申万行业名称']) else '其他行业'

        merged_df['申万行业名称'] = merged_df.apply(get_industry, axis=1)
        merged_df['股票代码'] = merged_df['股票代码'].astype(str).str.zfill(6)
        merged_df['企业名称'] = merged_df['企业名称'].fillna('未知企业').astype(str)
        merged_df['年份'] = merged_df['年份'].fillna(0).astype(int)

        for col in TECH_DIM_COLS:
            merged_df[col] = pd.to_numeric(merged_df[col], errors='coerce').fillna(0)

        merged_df['数字化转型指数'] = merged_df[TECH_DIM_COLS].mean(axis=1).round(4)
        merged_df['企业标识'] = merged_df.apply(lambda x: f"{x['股票代码']} | {x['企业名称']}", axis=1)
        merged_df = merged_df.dropna(subset=['企业标识', '年份']).reset_index(drop=True)

        return merged_df, f"✅ 数据加载完成！总记录数：{len(merged_df)} | 行业匹配率：{len(merged_df[merged_df['申万行业名称'] != '其他行业'])/len(merged_df):.2%}"

    except Exception as e:
        return None, f"❌ 数据加载失败：{str(e)}\n错误详情：{type(e).__name__}"

def generate_chart_figs(company_df, industry_df, selected_name, industry_name, year_start, year_end):
    """生成图表对象（不转图片）"""
    # 1. 总词频趋势图
    fig_total_freq = go.Figure()
    if not company_df.empty:
        fig_total_freq.add_trace(go.Scatter(
            x=company_df['年份'],
            y=company_df['总词频'],
            mode='lines+markers+text',
            name=f'{selected_name} 总词频',
            line=dict(color=COLOR_PALETTE['primary'], width=3),
            marker=dict(size=8),
            text=[f'{v}' for v in company_df['总词频']],
            textposition='top center'
        ))
    if not industry_df.empty:
        fig_total_freq.add_trace(go.Scatter(
            x=industry_df['年份'],
            y=industry_df['总词频'],
            mode='lines+markers',
            name=f'{industry_name} 行业平均词频',
            line=dict(color=COLOR_PALETTE['secondary'], width=3, dash='dash'),
            marker=dict(size=8)
        ))
    fig_total_freq.update_layout(
        title=f'{selected_name} 总词频趋势（{year_start}-{year_end}）',
        xaxis_title='年份',
        yaxis_title='总词频',
        template='plotly_white',
        height=500,
        legend=dict(orientation="h", yanchor="bottom", y=-0.2)
    )

    # 2. 行业对比折线图
    fig_industry = go.Figure()
    if not company_df.empty:
        fig_industry.add_trace(go.Scatter(
            x=company_df['年份'],
            y=company_df['数字化转型指数'],
            mode='lines+markers+text',
            name=f'{selected_name} 转型指数',
            line=dict(color=COLOR_PALETTE['primary'], width=4),
            marker=dict(size=10),
            text=[f'{v:.2f}' for v in company_df['数字化转型指数']],
            textposition='top center'
        ))
    if not industry_df.empty:
        fig_industry.add_trace(go.Scatter(
            x=industry_df['年份'],
            y=industry_df['数字化转型指数'],
            mode='lines+markers',
            name=f'{industry_name} 行业平均指数',
            line=dict(color=COLOR_PALETTE['secondary'], width=3, dash='dash'),
            marker=dict(size=8)
        ))
    fig_industry.update_layout(
        title=f'{selected_name} vs 行业转型指数对比',
        xaxis_title='年份',
        yaxis_title='数字化转型指数',
        template='plotly_white',
        height=500,
        legend=dict(orientation="h", yanchor="bottom", y=-0.2)
    )

    return {'total_freq': fig_total_freq, 'industry_compare': fig_industry}

def generate_report(company_df, selected_name, selected_code, industry_name, year_start, year_end):
    """生成无图片的Word报告（仅文本+表格）"""
    if company_df.empty:
        return None
    
    doc = Document()
    
    # 报告标题
    title = doc.add_heading('企业数字化转型分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息
    doc.add_heading('一、基本信息', level=1)
    info_table = doc.add_table(rows=1, cols=2)
    info_table.style = 'Table Grid'
    hdr_cells = info_table.rows[0].cells
    hdr_cells[0].text = '项目'
    hdr_cells[1].text = '内容'
    
    # 填充基本信息
    info_rows = [
        ('企业名称', selected_name),
        ('股票代码', selected_code),
        ('所属行业', industry_name),
        ('分析时间范围', f'{year_start}-{year_end}'),
        ('数据记录数', f'{len(company_df)}条'),
        ('平均总词频', f'{company_df["总词频"].mean():.2f}'),
        ('平均数字化转型指数', f'{company_df["数字化转型指数"].mean():.4f}'),
        ('最高转型指数', f'{company_df["数字化转型指数"].max():.4f}'),
        ('最低转型指数', f'{company_df["数字化转型指数"].min():.4f}')
    ]
    
    for row in info_rows:
        row_cells = info_table.add_row().cells
        row_cells[0].text = row[0]
        row_cells[1].text = row[1]
    
    # 年度详细数据
    doc.add_heading('二、年度详细数据', level=1)
    data_cols = ['年份', '总词频', '数字化转型指数'] + TECH_DIM_COLS[:6]  # 仅展示前6个技术维度
    data_table = doc.add_table(rows=1, cols=len(data_cols))
    data_table.style = 'Table Grid'
    
    # 数据表格表头
    hdr_cells = data_table.rows[0].cells
    for idx, col in enumerate(data_cols):
        hdr_cells[idx].text = col
    
    # 填充年度数据
    for _, row in company_df.iterrows():
        row_cells = data_table.add_row().cells
        for idx, col in enumerate(data_cols):
            val = row[col]
            if isinstance(val, float):
                row_cells[idx].text = f'{val:.4f}' if col == '数字化转型指数' else f'{val:.2f}'
            else:
                row_cells[idx].text = str(val)
    
    # 技术维度分析
    doc.add_heading('三、核心技术维度分析', level=1)
    tech_summary = doc.add_paragraph()
    tech_summary.add_run('1. 最高词频技术维度：').bold = True
    top_tech = company_df[TECH_DIM_COLS].mean().idxmax()
    tech_summary.add_run(f'{top_tech}（平均词频：{company_df[top_tech].mean():.2f}）\n')
    
    tech_summary.add_run('2. 最低词频技术维度：').bold = True
    low_tech = company_df[TECH_DIM_COLS].mean().idxmin()
    tech_summary.add_run(f'{low_tech}（平均词频：{company_df[low_tech].mean():.2f}）\n')
    
    # 趋势分析
    doc.add_heading('四、趋势分析', level=1)
    trend_para = doc.add_paragraph()
    trend_para.add_run('1. 总词频趋势：').bold = True
    if company_df['总词频'].iloc[-1] > company_df['总词频'].iloc[0]:
        trend_para.add_run('呈上升趋势，增长幅度：')
        trend_para.add_run(f'{((company_df["总词频"].iloc[-1] / company_df["总词频"].iloc[0]) - 1) * 100:.2f}%').bold = True
    else:
        trend_para.add_run('呈下降趋势，下降幅度：')
        trend_para.add_run(f'{(1 - (company_df["总词频"].iloc[-1] / company_df["总词频"].iloc[0])) * 100:.2f}%').bold = True
    
    trend_para.add_run('\n2. 数字化转型指数趋势：').bold = True
    if company_df['数字化转型指数'].iloc[-1] > company_df['数字化转型指数'].iloc[0]:
        trend_para.add_run('呈上升趋势，企业数字化转型进程加快')
    else:
        trend_para.add_run('呈下降趋势，需关注数字化转型投入')
    
    # 生成时间
    doc.add_paragraph(f'\n报告生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    
    # 保存到字节流
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# -------------------------- 自动加载数据 --------------------------
if not st.session_state.data_loaded:
    with st.spinner("🔄 加载数据中..."):
        data, msg = load_data()
        if data is not None:
            st.session_state.merged_data = data
            st.session_state.data_loaded = True
            st.success(msg)
        else:
            st.error(msg)

# -------------------------- 主界面 --------------------------
st.title("📊 企业数字化转型指数分析平台")

if st.session_state.data_loaded:
    df = st.session_state.merged_data

    # 企业筛选
    st.subheader("🔍 企业筛选")
    col1, col2, col3 = st.columns([2, 1, 1])
    with col1:
        company_options = sorted(df['企业标识'].unique())[:100]
        selected_company = st.selectbox("选择企业", company_options, index=0 if len(company_options) > 0 else None)
        selected_code = selected_company.split(' | ')[0] if selected_company else '000000'
        selected_name = selected_company.split(' | ')[1] if selected_company else '未知企业'
    with col2:
        valid_years = sorted(df['年份'].unique())
        year_start = st.selectbox("起始年份", valid_years, index=0 if len(valid_years) > 0 else None)
    with col3:
        year_end = st.selectbox("结束年份", valid_years, index=len(valid_years) - 1 if len(valid_years) > 0 else None)

    # 筛选数据
    company_df = df[
        (df['股票代码'] == selected_code) &
        (df['年份'] >= year_start) &
        (df['年份'] <= year_end)
    ].sort_values('年份').reset_index(drop=True)

    industry_name = company_df['申万行业名称'].iloc[0] if not company_df.empty else '其他行业'
    industry_df = df[
        (df['申万行业名称'] == industry_name) &
        (df['年份'] >= year_start) &
        (df['年份'] <= year_end)
    ].groupby('年份').agg({'总词频': 'mean', '数字化转型指数': 'mean'}).reset_index()

    # 生成图表
    if not company_df.empty:
        st.session_state.chart_figs = generate_chart_figs(
            company_df, industry_df, selected_name, industry_name, year_start, year_end
        )

    # 核心指标
    st.subheader("📋 核心指标")
    if not company_df.empty:
        col1, col2, col3, col4 = st.columns(4)
        with col1: st.metric("企业名称", selected_name)
        with col2: st.metric("股票代码", selected_code)
        with col3: st.metric("所属行业", industry_name)
        with col4: st.metric("平均转型指数", f"{company_df['数字化转型指数'].mean():.4f}")

    # 词频趋势分析
    st.subheader("📈 词频趋势分析")
    tab1, tab2 = st.tabs(["总词频趋势", "技术维度词频趋势"])
    with tab1:
        if 'total_freq' in st.session_state.chart_figs:
            st.plotly_chart(st.session_state.chart_figs['total_freq'], use_container_width=True)
        else:
            st.info("暂无足够数据生成趋势图")
    with tab2:
        selected_tech = st.multiselect("选择技术维度", TECH_DIM_COLS, default=TECH_DIM_COLS[:4], key='tech_dim_select')
        if selected_tech and not company_df.empty:
            fig_tech_freq = go.Figure()
            for idx, tech in enumerate(selected_tech):
                fig_tech_freq.add_trace(go.Scatter(
                    x=company_df['年份'], y=company_df[tech], mode='lines+markers', name=tech,
                    line=dict(color=list(COLOR_PALETTE.values())[idx % len(COLOR_PALETTE)], width=2)
                ))
            fig_tech_freq.update_layout(title=f'{selected_name} 技术维度词频趋势', xaxis_title='年份', yaxis_title='词频', template='plotly_white', height=500)
            st.plotly_chart(fig_tech_freq, use_container_width=True)

    # 行业对比分析
    st.subheader("🏭 行业对比分析")
    if 'industry_compare' in st.session_state.chart_figs:
        fig_industry = st.session_state.chart_figs['industry_compare']
        industry_names = [str(name) for name in df[df['申万行业名称'] != '其他行业']['申万行业名称'].unique() if name.strip()][:20]
        other_industries = st.multiselect("添加其他行业对比", industry_names, default=[], key='other_industry')
        color_idx = 2
        for ind in other_industries:
            ind_data = df[(df['申万行业名称'] == ind) & (df['年份'] >= year_start) & (df['年份'] <= year_end)].groupby('年份')['数字化转型指数'].mean().reset_index()
            if not ind_data.empty:
                fig_industry.add_trace(go.Scatter(
                    x=ind_data['年份'], y=ind_data['数字化转型指数'], mode='lines+markers', name=f'{ind} 行业平均',
                    line=dict(color=list(COLOR_PALETTE.values())[color_idx % len(COLOR_PALETTE)], width=2), marker=dict(size=6)
                ))
                color_idx += 1
        st.plotly_chart(fig_industry, use_container_width=True)
    else:
        st.info("暂无足够数据生成行业对比图")

    # 详细数据
    st.subheader("📝 详细数据")
    if not company_df.empty:
        display_cols = ['年份', '股票代码', '企业名称', '申万行业名称', '总词频', '数字化转型指数'] + TECH_DIM_COLS
        st.dataframe(
            company_df[display_cols], use_container_width=True, hide_index=True,
            column_config={"数字化转型指数": st.column_config.NumberColumn(format="%.4f"), "总词频": st.column_config.NumberColumn(format="%d")},
            height=300
        )

    # 数据下载
    st.subheader("💾 数据下载")
    if not company_df.empty:
        col1, col2 = st.columns(2)
        with col1:
            csv_data = company_df[display_cols].to_csv(index=False, encoding='utf-8-sig')
            st.download_button("下载CSV数据", data=csv_data, file_name=f"{selected_name}_{year_start}-{year_end}_转型数据.csv", use_container_width=True)
        with col2:
            excel_buffer = io.BytesIO()
            with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                company_df[display_cols].to_excel(writer, sheet_name='转型数据', index=False)
            st.download_button("下载Excel数据", data=excel_buffer, file_name=f"{selected_name}_{year_start}-{year_end}_转型数据.xlsx", use_container_width=True)

else:
    st.info("💡 数据加载中，请稍候...")

# -------------------------- 侧边栏（恢复报告下载） --------------------------
with st.sidebar:
    st.header("📄 报告下载")
    
    # 报告生成按钮
    if st.button("生成分析报告", type="primary", use_container_width=True):
        if company_df.empty:
            st.error("❌ 暂无数据生成报告，请选择有效企业和时间范围")
        else:
            with st.spinner("📝 正在生成报告..."):
                report_buffer = generate_report(company_df, selected_name, selected_code, industry_name, year_start, year_end)
                if report_buffer:
                    st.download_button(
                        label="📥 下载Word报告",
                        data=report_buffer,
                        file_name=f"{datetime.now().strftime('%Y%m%d')}_{selected_name}_数字化转型分析报告.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                    st.success("✅ 报告生成成功！点击按钮下载")
                else:
                    st.error("❌ 报告生成失败")

    st.divider()
    
    # 平台说明
    st.header("ℹ️ 平台说明")
    st.info(f"""
    📁 当前数据目录：
    {BASE_DIR}/年报下载

    📄 词频数据文件：
    {WORDFREQ_FILE}

    📊 行业数据文件：
    {INDUSTRY_FILE}
    """)

    st.divider()
    st.markdown("""
    📅 更新时间：2025年12月  
    🛠️ 技术栈：Streamlit + Plotly + Pandas  
    ⚡ 核心功能：词频趋势 + 行业对比 + 数据/报告下载
    🎯 行业匹配：股票代码优先 + 企业名称兜底
    🚀 适配说明：移除图片导出，兼容云端部署
    """)

# -------------------------- 页脚 --------------------------
st.divider()
st.markdown(f"© {datetime.now().year} 企业数字化转型分析平台 | 适配Streamlit Cloud部署")
